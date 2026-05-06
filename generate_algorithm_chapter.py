"""
生成 HAPPO-GNN-RL 论文算法章节 Word 文档
直接可放入学位论文第3章（模型建立）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ==================== 页面设置 ====================
section = doc.sections[0]
section.page_width = Inches(8.27)    # A4
section.page_height = Inches(11.69)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ==================== 字体设置 ====================
def set_run_font(run, size=12, bold=False, color=None, name='Times New Roman'):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_body(doc, text, indent=0, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_formula(doc, formula_text, caption=""):
    """添加居中的公式块（用等宽字体模拟公式效果）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(32)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 添加背景色块效果
    run = p.add_run(formula_text)
    run.font.size = Pt(11)
    run.font.name = 'Cambria Math'
    run.font.bold = False
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.line_spacing = Pt(18)
        cap_r = cap_p.add_run(caption)
        cap_r.font.size = Pt(10)
        cap_r.font.italic = True
        cap_r.font.name = 'Times New Roman'
        cap_r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        cap_r.font.color.rgb = RGBColor(80, 80, 80)

def add_algo_block(doc, algo_lines, title="", explain_lines=None):
    """添加算法框（表格模拟）"""
    # 算法标题
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 用表格模拟算法代码块
    table = doc.add_table(rows=len(algo_lines), cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格背景色（浅灰）
    tbl = table._tbl
    tblPr = tbl.tblPr
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    tblPr.append(shd)

    for i, line in enumerate(algo_lines):
        cell = table.cell(i, 0)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.line_spacing = Pt(18)

        if line.strip().startswith('//') or line.strip().startswith('#'):
            # 注释行
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.italic = True
            run.font.name = 'Consolas'
        else:
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(30, 30, 30)

    doc.add_paragraph()  # 空行

    # 中文通俗讲解
    if explain_lines:
        for exp in explain_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.line_spacing = Pt(22)
            # 缩进处理
            if exp.startswith('    ') or exp.startswith('\t'):
                p.paragraph_format.left_indent = Cm(1.5)
                run = p.add_run(exp.strip())
            else:
                run = p.add_run(exp)
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.color.rgb = RGBColor(50, 50, 50)

def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.italic = True
    return p

# ==================== 章节标题 ====================
add_heading(doc, '第 3 章  模型建立', level=1)
add_heading(doc, '3.1  问题分析', level=2)

add_body(doc,
    '电动汽车充电站的动态调度优化问题，本质上是一个多智能体序贯决策问题。在充电站网络中，'
    '各站点之间存在显著的空间相关性（相邻区域的站点负荷往往呈现同步波动），同时各站点内的'
    '不同车型（公交车、乘用车、两轮电动车）具有差异化的充电需求与优先级。'
    '本文提出 HAPPO-GNN-RL（Heterogeneous Attentive PPO with Graph Neural Network）模型，'
    '利用图注意力网络（GAT）建模站点间的空间依赖关系，采用异质多智能体近端策略优化（HAPPO）'
    '实现动态调度决策。模型的核心在于：先通过物理先验（互信息和格兰杰因果）构建站点关联强度，'
    '再通过 GAT 的动态注意力机制学习实时的空间关系，最终在 HAPPO 框架下完成多车型差异化调度。')

# ==================== 3.2 物理先验构建 ====================
add_heading(doc, '3.2  物理先验：站点关联强度矩阵', level=2)

add_body(doc,
    '在构建图注意力网络之前，需要预先建模充电站之间的物理关联关系。本文采用互信息（Mutual Information, MI）'
    '与格兰杰因果检验（Granger Causality）相结合的方式，计算站点之间的固定物理关联强度 $S_{ij}$。'
    '互信息衡量两个站点负荷序列的统计相关性，格兰杰因果检验则判断一个站点的历史数据能否有效预测另一站点的未来负荷。'
    '二者加权求和，构成站点对的物理关联强度。')

add_formula(doc,
    r'$S_{ij} = \text{MI}(i,j) + \text{Granger}(i \to j)$',
    '式（3.1）  站点关联强度矩阵计算公式')

add_body(doc,
    '其中：$\\text{MI}(i,j)$ 表示站点 $i$ 与站点 $j$ 负荷序列的互信息，'
    '$\\text{Granger}(i \\to j)$ 表示站点 $i$ 对站点 $j$ 的格兰杰因果强度。'
    '该关联强度矩阵 $S \\in \\mathbb{R}^{N \\times N}$（$N$ 为站点数）作为物理先验，'
    '为后续注意力融合提供方向性引导。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• MI（互信息）：看两个站点"历史上关系强不强"——如果两个站同时高峰、同时低谷，说明它们有内在联系。'
    '  计算方法：将两个站点的负荷离散化，统计联合分布与边缘分布，代入 $MI(X,Y) = \\sum \\sum p(x,y) \\log\\frac{p(x,y)}{p(x)p(y)}$。')
add_body(doc,
    '• Granger（格兰杰因果）：看一个站的历史数据能不能预测另一个站的未来——如果 A 站先涨，B 站也跟着涨，'
    '  说明 A 对 B 有因果影响。计算方法：构建 VAR 模型，用 F 检验判断 $i$ 的滞后项对 $j$ 是否有显著预测能力。'
    '  有则 $S_{ij}$ 增大，无则不变。')

add_body(doc,
    '两个指标加在一起，得到每对站点之间的物理关系分数。这个分数不是模型学出来的，而是根据领域知识预先算好的，'
    '它告诉注意力机制："这几个站本来就有关系，先给他们加点初始权重。"'
    '后续 GAT 的动态注意力会在这个基础上进一步调整。')

# ==================== 3.3 GAT 特征提取 ====================
add_heading(doc, '3.3  图注意力网络（GAT）：空间特征提取', level=2)

add_body(doc,
    '充电站网络本质上是一个非欧几里得结构的图数据。为有效建模站点间的动态空间依赖关系，'
    '本文采用图注意力网络（Graph Attention Network, GAT）对充电站网络进行特征提取。'
    'GAT 通过注意力机制自适应地为每个站点邻居分配权重，避免了传统 GCN 固定卷积核的局限性。')

# ----- 动态注意力 -----
add_heading(doc, '3.3.1  动态注意力分数计算', level=3)

add_body(doc,
    '给定站点 $i$ 和邻居站点 $j$ 的节点特征向量 $h_i, h_j \\in \\mathbb{R}^d$，'
    'GAT 首先通过一个可学习的线性变换将特征映射到注意力空间，然后计算两站的注意力分数：')

add_formula(doc,
    r'$\omega_{ij} = \text{LeakyReLU}\Big( W_a \cdot \big[ h_i \oplus h_j \big] \Big)$',
    '式（3.2）  动态注意力分数计算')

add_body(doc,
    '其中：$\\oplus$ 表示向量拼接（concatenation），$W_a \\in \\mathbb{R}^{2d \\times 1}$ 为可学习权重矩阵，'
    '$\\text{LeakyReLU}$ 保证注意力分数非负。'
    '该注意力分数 $\\omega_{ij}$ 由 GAT 在训练过程中自主学习，代表站点 $j$ 对站点 $i$ 当前时刻的实时影响程度。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '① 拼接（$\\oplus$）：把站点 i 的特征和站点 j 的特征拼成一个长向量。比如每个站点特征是64维，'
    '拼起来就是128维。这一步是为了让模型同时看到"我自己是什么"和"邻居是什么"。')
add_body(doc,
    '② 线性压缩（$W_a$）：把128维向量压缩成1个数字，就是注意力分数。这个 $W_a$ 是模型学出来的，不是人工设定的。'
    '模型会自己决定哪些特征重要、哪些特征不重要。')
add_body(doc,
    '③ LeakyReLU：保证分数不为负。如果分数是负的，就变成0。这相当于一个"门"，太负面的关系直接忽略。')

# ----- 注意力融合 -----
add_heading(doc, '3.3.2  注意力融合：物理先验 + 动态学习', level=3)

add_body(doc,
    '单纯的动态注意力完全依赖数据驱动，缺乏领域知识的引导；'
    '单纯的物理先验无法捕捉实时变化的负荷波动。'
    '本文将动态注意力 $\\omega_{ij}$ 与物理先验 $S_{ij}$ 进行加权融合，得到最终的注意力原始分数：')

add_formula(doc,
    r'$e_{ij} = \omega_{ij} + \gamma \cdot S_{ij}$',
    '式（3.3）  注意力融合')

add_body(doc,
    '其中 $\\gamma$ 为物理先验系数，由模型在学习过程中自适应调整。'
    '当 $\\gamma > 0$ 时，物理先验对注意力分配产生正向引导；当 $\\gamma = 0$ 时，退化为纯动态注意力。'
    '融合后的分数 $e_{ij}$ 再通过 $\\text{Softmax}$ 归一化：')

add_formula(doc,
    r'$\alpha_{ij} = \text{Softmax}\Big( e_{ij} \Big) = \frac{\exp(e_{ij})}{\sum_{k \in \mathcal{N}_i} \exp(e_{ik})}$',
    '式（3.4）  注意力权重归一化')

add_body(doc,
    '其中 $\\mathcal{N}_i$ 为站点 $i$ 的邻居节点集合，$\\alpha_{ij}$ 为站点 $j$ 对站点 $i$ 的最终注意力权重，'
    '满足 $\\sum_{j \\in \\mathcal{N}_i} \\alpha_{ij} = 1$。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• 融合的意思：把"站点间固定的历史关系"（物理先验 $S_{ij}$）和'
    '"模型此时此刻自己学到的关系"（动态注意力 $\\omega_{ij}$）加在一起。')
add_body(doc,
    '• $\\gamma$ 的作用：如果模型发现物理先验很准，$\\gamma$ 就会自动变大；如果物理先验不准（和实际数据冲突），'
    '$\\gamma$ 就会变小。模型在训练中自己学会平衡这两者。')
add_body(doc,
    '• Softmax 的作用：把注意力分数归一化到0~1之间，并且所有邻居的权重加起来等于1。'
    '这使得注意力权重可以直接作为加权平均的系数。')

# ----- 图特征聚合 -----
add_heading(doc, '3.3.3  图特征聚合与更新', level=3)

add_body(doc,
    '获得注意力权重后，GAT 通过加权聚合机制更新每个站点的特征表示：')

add_formula(doc,
    r'$h_i^{\prime(l)} = \sum_{j \in \mathcal{N}_i} \alpha_{ij}^{(l)} \cdot W_V^{(l)} h_j^{(l)}$',
    '式（3.5）  图特征聚合')

add_body(doc,
    '其中 $W_V^{(l)} \\in \\mathbb{R}^{d \\times d}$ 为第 $l$ 层 GAT 的值矩阵，'
    '$h_j^{(l)}$ 为邻居站点 $j$ 在第 $l$ 层的特征，'
    '$\\alpha_{ij}^{(l)}$ 为对应的注意力权重。'
    '本文采用多头注意力（Multi-Head Attention）机制，将 $K$ 个注意力头的输出拼接：')

add_formula(doc,
    r'$h_i^{(l+1)} = \bigoplus_{k=1}^{K} \sigma\Big( \sum_{j \in \mathcal{N}_i} \alpha_{ij}^{(k,l)} W_V^{(k,l)} h_j^{(l)} \Big)$',
    '式（3.6）  多头注意力聚合')

add_body(doc,
    '其中 $\\sigma$ 为 $\\text{ELU}$ 激活函数，$K$ 为注意力头数（本文取 $K=4$）。'
    '为缓解深层 GAT 的梯度消失问题，本文引入残差连接（Residual Connection）：')

add_formula(doc,
    r'$h_i^{(l+1)} = \sigma\Big( h_i^{\prime(l)} \Big) + h_i^{(l)}$',
    '式（3.7）  残差连接')

add_body(doc,
    '残差连接保证即使网络很深，原始特征信息也不会完全丢失，模型训练更加稳定。'
    '经过 $L$ 层 GAT 后，每个站点 $i$ 获得一个融合了全局空间信息的特征向量 $h_i^{(L)}$。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• 特征聚合（式3.5）：把每个邻居的特征按注意力权重加权求和。'
    '比如 A 站注意力权重 0.6，B 站 0.3，C 站 0.1，则 A 站的新特征 = 0.6×A + 0.3×B + 0.1×C。'
    '这相当于每个站点在"询问"邻居："根据你们的状态，我应该怎么调整？"')
add_body(doc,
    '• 多头注意力（式3.6）：用4个不同的"眼睛"同时看问题，每个头学一种不同的注意力模式。'
    '比如头1关注电价相关、头2关注负荷相关、头3关注绿电相关、头4关注队列相关。最后拼接4个头的输出，'
    '得到更丰富、更鲁棒的特征表示。')
add_body(doc,
    '• 残差连接（式3.7）：直接把上一层的输入加到输出上。'
    '这就像"高速公路"——即使中间层学坏了，原始信息还能直接传过去，不会消失。'
    '有了残差连接，4层甚至更深的 GAT 都能稳定训练。')

# ==================== 算法1：GAT特征提取 ====================
add_heading(doc, '算法 3.1  基于物理先验的 GAT 空间特征提取算法', level=2)

algo1 = [
    '# 输入：站点数 N，特征维度 d，层数 L，注意力头数 K',
    '# 输出：融合空间特征的节点表示 H^{(L)} = {h_1^{(L)}, ..., h_N^{(L)}}',
    '',
    '# Step 1：构建物理先验关联强度矩阵 S',
    'for each station pair (i, j) in N × N:',
    '    MI(i,j) ← 计算站点 i 与 j 负荷序列的互信息',
    '    Granger(i→j) ← 格兰杰因果检验（i 对 j 的因果强度）',
    '    S_{ij} ← MI(i,j) + Granger(i→j)          # 式（3.1）',
    'end for',
    '',
    '# Step 2：GAT 多层特征提取',
    'H^{(0)} ← 初始化站点特征矩阵                    # 原始特征向量拼接',
    'for l = 1 to L:                                 # 遍历 GAT 层',
    '    for each head k = 1 to K:                   # 遍历注意力头',
    '        for each station pair (i, j):',
    '            # ① 动态注意力分数',
    '            e_{ij}^{(k,l)} ← LeakyReLU(W_a^{(k,l)} · [h_i^{(l-1)} ⊕ h_j^{(l-1)}])    # 式（3.2）',
    '            # ② 融合物理先验',
    '            e_{ij}^{(k,l)} ← e_{ij}^{(k,l)} + γ · S_{ij}                            # 式（3.3）',
    '        end for',
    '        # ③ Softmax 归一化',
    '        α_{ij}^{(k,l)} ← exp(e_{ij}^{(k,l)}) / Σ_{k∈N_i} exp(e_{ik}^{(k,l)})        # 式（3.4）',
    '        # ④ 邻居特征聚合',
    '        h_i^{(k,l)} ← Σ_{j∈N_i} α_{ij}^{(k,l)} · W_V^{(k,l)} · h_j^{(l-1)}          # 式（3.5）',
    '    end for',
    '    # ⑤ 多头拼接 + 残差连接',
    '    h_i^{(l)} ← (⊕_{k=1}^{K} σ(h_i^{(k,l)})) + h_i^{(l-1)}                           # 式（3.6）+（3.7）',
    'end for',
    '',
    '# Step 3：输出空间感知特征',
    'return H^{(L)} = {h_1^{(L)}, ..., h_N^{(L)}}',
]

algo1_explain = [
    '【通俗讲解】',
    '    ① Step 1：预先算好站点间的历史关联强度（物理先验），这一步不用模型训练，直接用数据计算。',
    '    ② Step 2 第1-4行：对每对站点计算"此时此刻有多相关"。先把特征拼接 → 压缩成1个分数 → LeakyReLU → 加物理先验。',
    '    ③ Step 2 第5行：Softmax 把分数变成可解释的权重（总和=1）。',
    '    ④ Step 2 第6行：按注意力权重把邻居特征加权求和，这一步让每个站吸收周围站点的信息。',
    '    ⑤ Step 2 第8行：4个注意力头并行学习，最后拼接；同时加残差连接防止训练崩溃。',
    '    ⑥ 重复 L 层：信息在站点间传播得更远（1层看邻居，2层看邻居的邻居，……）。',
]

add_algo_block(doc, algo1, title="", explain_lines=algo1_explain)

# ==================== 3.4 HAPPO 策略网络 ====================
add_heading(doc, '3.4  HAPPO 策略网络与动作选择', level=2)

add_body(doc,
    '经过 GAT 特征提取后，每个站点 $i$ 拥有融合了全局空间信息的特征表示 $h_i = h_i^{(L)}$。'
    '本文采用异质多智能体近端策略优化（HAPPO）的框架，分别针对三种车型（公交车、乘用车、两轮电动车）'
    '建立独立的策略网络，实现差异化的调度决策。')

# ----- 策略网络 MLP -----
add_heading(doc, '3.4.1  策略网络结构', level=3)

add_body(doc,
    '策略网络采用两层全连接 MLP，将站点特征 $h_i$ 映射为动作空间中的评分向量：')

add_formula(doc,
    r'$z_1 = W_1 \cdot h_i + b_1,\ W_1 \in \mathbb{R}^{128 \times 64} \\ '
    r'z_2 = \text{LeakyReLU}(z_1) \\ '
    r'z_3 = W_2 \cdot z_2 + b_2,\ W_2 \in \mathbb{R}^{A \times 128} \\ '
    r'\pi_\theta(a_i | s_i) = \text{Softmax}(z_3)$',
    '式（3.8）  策略网络结构')

add_body(doc,
    '其中：$W_1, W_2$ 为可学习权重，$A$ 为动作空间维度（包含充电功率档位分配、排队优先级等），'
    '$\\pi_\\theta(a_i | s_i)$ 为在状态 $s_i$ 下选择动作 $a_i$ 的概率分布。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• 第1层（z1）：把64维站点特征扩展到128维，增加模型容量，让特征之间的关系有空间表达。')
add_body(doc,
    '• LeakyReLU：激活函数，引入非线性。普通 ReLU 把负数全变成0，LeakyReLU 让负数乘一个小系数（0.01）保留一点点信息，'
    '防止"死亡神经元"问题。')
add_body(doc,
    '• 第2层（z3）：从128维压缩到动作数A维。每个维度代表一个动作的"原始分数"（logit）。')
add_body(doc,
    '• Softmax：把分数归一化为概率。分数最高的动作被选中的概率最大，但其他动作也有非零概率（保证探索）。')

# ----- 动作掩码 -----
add_heading(doc, '3.4.2  异质动作掩码机制', level=3)

add_body(doc,
    '不同车型的充电站，其合法动作集合不同。例如，公交车充电站的可用功率档位与两轮电动车充电站不同，'
    '某些时段部分充电桩可能处于维护状态。本文引入动作掩码（Action Mask）机制，'
    '对非法动作进行强制屏蔽：')

add_formula(doc,
    r'\pi_{\text{masked}}(a_i | s_i) = \pi_\theta(a_i | s_i) \odot M(s_i)',
    '式（3.9）  动作掩码')

add_formula(doc,
    r'\pi_{\text{norm}}(a_i | s_i) = \frac{\pi_{\text{masked}}(a_i | s_i)}{\sum_{a^\prime} \pi_{\text{masked}}(a^\prime | s_i)}',
    '式（3.10）  掩码后概率重归一化')

add_body(doc,
    '其中 $M(s_i) \\in \\{0,1\\}^A$ 为动作掩码向量，1表示该动作合法可用，0表示非法或不可用；'
    '$\\odot$ 为按位乘法；$\\pi_{\\text{norm}}$ 为掩码后的合法动作概率分布，保证其和为1。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• 比如有5个动作，掩码 M = [1, 1, 0, 1, 0]，意思是动作3和动作5在当前状态下不可用（如充电桩已满或处于维护状态）。')
add_body(doc,
    '• $\\odot$（按位乘）：把不可用动作的原始概率直接变成0。比如原始 [0.2, 0.3, 0.1, 0.3, 0.1] × [1,1,0,1,0] = [0.2, 0.3, 0, 0.3, 0]。')
add_body(doc,
    '• 重归一化：原来5个动作和为1，屏蔽后和变成0.8。式（3.10）按比例放大，让剩下的合法动作概率重新加起来等于1，'
    '保证模型不会因为掩码而"不知所措"。')

# ----- 最终动作选择 -----
add_heading(doc, '3.4.3  最终调度决策', level=3)

add_body(doc,
    '掩码重归一化后，选择概率最大的合法动作作为当前调度决策：')

add_formula(doc,
    r'a_i^* = \arg\max_{a_i} \; \pi_{\text{norm}}(a_i | s_i)',
    '式（3.11）  最终调度动作选择')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• 选出掩码后概率最大的动作。比如 [0.33, 0.50, 0, 0.17, 0] → 最大概率是0.50，对应动作2，'
    '所以最终选择"动作2"（如"分配到3号充电桩，功率为7kW"）。')
add_body(doc,
    '• 注意：这是确定性选择（argmax），训练时用概率采样（按概率大小随机抽），'
    '推理时用确定性选择。这叫"训练-推理不对称性"，是 RL 中的标准做法。')

# ==================== 算法2：策略网络 ====================
add_heading(doc, '算法 3.2  掩码策略网络与动作选择算法', level=2)

algo2 = [
    '# 输入：站点特征 h_i ∈ R^d，动作掩码 M(s_i) ∈ {0,1}^A',
    '# 输出：掩码归一化动作概率分布 π_norm，动作 a_i^*',
    '',
    '# ── 策略网络前向传播 ──',
    'z_1 ← W_1 · h_i + b_1                              # W_1 ∈ R^{128×64}',
    'z_2 ← LeakyReLU(z_1)                               # 激活，引入非线性',
    'z_3 ← W_2 · z_2 + b_2                              # W_2 ∈ R^{A×128}',
    'π_θ ← Softmax(z_3)                                 # 原始动作概率，式（3.8）',
    '',
    '# ── 动作掩码 ──',
    'π_masked ← π_θ ⊙ M(s_i)                            # 非法动作概率置0，式（3.9）',
    "sum_masked ← Σ_{a'} π_masked[a'']                  # 合法动作概率之和",
    'π_norm[a] ← π_masked[a] / sum_masked   for all a  # 重新归一化，式（3.10）',
    '',
    '# ── 动作选择 ──',
    'a_i^* ← argmax_a π_norm[a]                         # 选择概率最大动作，式（3.11）',
    'return π_norm, a_i^*',
]

algo2_explain = [
    '【通俗讲解】',
    '    ① 前向传播：特征 → 线性变换1 → LeakyReLU → 线性变换2 → Softmax，得到每个动作的原始概率。',
    '    ② 掩码处理：把当前不可用的动作概率直接清零（乘以0）。',
    '    ③ 归一化：把剩余合法动作的概率按比例放大，确保加起来还是1。这样模型总在做"合法动作中的最优"。',
    '    ④ 动作选择：选概率最大的那个。训练时概率采样（探索），推理时直接取最大（利用）。',
    '    关键：掩码机制保证了"不同车型 × 不同站点 × 不同时段"的异质性，不会选出不存在的动作。',
]

add_algo_block(doc, algo2, title="", explain_lines=algo2_explain)

# ==================== 3.5 奖励函数 ====================
add_heading(doc, '3.5  奖励函数设计', level=2)

add_body(doc,
    '奖励函数定义了调度策略的优化目标。本文将充电站调度的核心目标建模为多目标加权奖励：'
    '最小化用户等待时间、降低充电成本、减少电网负荷峰谷差，同时最大化绿电利用。'
    '综合奖励函数定义为：')

add_formula(doc,
    r'R_t = - \Big( w_1 \cdot T_t + w_2 \cdot P_t + w_3 \cdot L_t - w_4 \cdot G_t \Big)',
    '式（3.12）  综合奖励函数')

add_body(doc,
    '其中 $T_t$ 为等待时间，$P_t$ 为充电成本，$L_t$ 为负荷不均衡度，$G_t$ 为绿电利用量；'
    '$w_1, w_2, w_3, w_4 > 0$ 为自适应权重系数，在训练过程中根据实时状态动态调整。'
    '具体而言：')

add_body(doc,
    '• $w_1(t) = f(\\text{queue\\_length}, \\text{waiting\\_time})$：队列越长、等待时间越大，$w_1$ 越大，'
    '优先优化等待时间。')
add_body(doc,
    '• $w_2(t) = g(\\text{grid\\_load}, \\text{transformer\\_capacity})$：电网负荷接近变压器容量上限时，'
    '$w_2$ 增大，降低负荷压力。')
add_body(doc,
    '• $w_3(t) = h(\\text{pv\\_output}, \\text{wind\\_output})$：绿电充足时（如光伏大发时段），'
    '$w_4$ 增大，优先消纳可再生能源。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• 负号（−）：本质上是最小化成本函数，所以取负变成最大化奖励函数，这是 RL 的标准做法。'
    '如果等待时间短、成本低、绿电用得多 → 括号里的值小 → 负号后变成大奖励。')
add_body(doc,
    '• 自适应权重：不是固定的权重，而是根据当前系统状态实时调整。比如早高峰公交车多，'
    '$w_1$ 就变大；中午光伏强，$w_4$ 就变大。模型自动学会"什么情况优先解决什么问题"。')
add_body(doc,
    '• 绿电利用（$G_t$）前是减号：绿电用得越多，括号里减去的越多，奖励越高。这激励模型优先在绿电充足时安排充电。')

# ==================== 3.6 HAPPO 训练算法 ====================
add_heading(doc, '3.6  HAPPO 多智能体训练算法', level=2)

add_body(doc,
    '本文采用异质多智能体近端策略优化（HAPPO）作为核心训练算法。'
    '与标准 MAPPO 不同，HAPPO 针对不同车型的异质性，设计了顺序更新（Sequential Update）机制：'
    '先更新公交智能体（优先级最高），再更新乘用车智能体，最后更新两轮车智能体。'
    '每个智能体使用 PPO 的裁剪损失函数进行优化：')

add_formula(doc,
    r'L^{\text{CLIP}}(\theta) = \mathbb{E}_t\left[ \min\Big( r_t(\theta) \cdot A_t,\ \text{clip}\big(r_t(\theta), 1-\epsilon, 1+\epsilon\big) \cdot A_t \Big) \right]',
    '式（3.13）  PPO 裁剪损失函数')

add_body(doc,
    '其中 $r_t(\\theta) = \\frac{\\pi_{\\theta}(a_t | s_t)}{\\pi_{\\theta_{\\text{old}}}(a_t | s_t)}$ '
    '为新策略与旧策略的概率比（重要性采样比），'
    '$A_t$ 为优势函数（衡量当前动作相对于平均水平的优劣），'
    '$\\epsilon = 0.2$ 为裁剪超参数，'
    '$\\text{clip}(r_t, 1-\\epsilon, 1+\\epsilon)$ 将策略更新幅度限制在 $[1-\\epsilon, 1+\\epsilon]$ 区间内，'
    '防止一次性更新过猛导致策略崩溃。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• $r_t$（重要性采样比）：新策略选择这个动作的概率 ÷ 旧策略选择这个动作的概率。'
    '如果 $r_t > 1$，说明新策略更倾向于选这个动作；如果 $r_t < 1$，说明新策略更不倾向于选它。')
add_body(doc,
    '• $A_t$（优势函数）：衡量这个动作"比平均水平好多少"。如果 $A_t > 0$，说明这个动作值得鼓励；'
    '如果 $A_t < 0$，说明应该减少选这个动作的概率。')
add_body(doc,
    '• clip（裁剪）：把 $r_t$ 限制在 [0.8, 1.2] 区间内。'
    '即使新策略强烈想做某个动作（$r_t=5$），也只按1.2倍更新，防止一步走太远。反之，即使强烈不想做（$r_t=0.1$），'
    '也只按0.8倍减少。这就像"步幅限制器"，保证训练稳定。')
add_body(doc,
    '• min：取 CLIP 损失与原始损失的最小值。当 $A_t > 0$（好动作）时，'
    '$r_t$ 增大的损失被限制，不让策略跑太远；当 $A_t < 0$（坏动作）时同理。')
add_body(doc,
    '• $\\mathbb{E}_t$：对一批数据（多个时间步、多个智能体）求平均估计，让梯度估计更稳定、方差更低。')

# ----- 梯度更新 -----
add_heading(doc, '3.6.1  梯度更新与策略优化', level=3)

add_body(doc,
    '策略网络的参数通过梯度上升更新（等价为梯度下降最小化负的裁剪损失）：')

add_formula(doc,
    r'\theta_{\text{new}} \leftarrow \theta_{\text{old}} - \eta \cdot \nabla_\theta L^{\text{CLIP}}(\theta)',
    '式（3.14）  策略网络梯度更新')

add_body(doc,
    '其中 $\\eta$ 为学习率，$\\nabla_\\theta$ 为损失函数对策略参数 $\\theta$ 的梯度。'
    '本文采用 Adam 优化器，并根据梯度范数进行自适应学习率调整。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• 梯度方向：指向损失函数增长最快的方向。减去梯度 = 沿损失减少最快的方向走一步。'
    '这就是标准的"梯度下降"，只不过 RL 里最小化损失 = 最大化期望累积奖励。')
add_body(doc,
    '• Adam 优化器：比普通 SGD 更聪明，会自动调整每层学习率。如果某参数梯度一直很小，Adam 增大它的学习率；'
    '如果梯度一直很大，Adam 减小它的学习率。训练更稳定、收敛更快。')
add_body(doc,
    '• 迭代更新：每轮训练结束后，用新的 $\\theta_{\\text{new}}$ 替换 $\\theta_{\\text{old}}$，'
    '然后重新收集数据、计算优势函数、再训练。循环往复，直到策略收敛。')

# ----- HAPPO 顺序更新 -----
add_heading(doc, '3.6.2  异质顺序更新机制', level=3)

add_body(doc,
    'HAPPO 的核心创新在于异质顺序更新（Heterogeneous Sequential Update）。'
    '不同车型的充电需求、功率约束和时间敏感性差异显著，采用统一更新方式会导致低优先级车型被忽视。'
    'HAPPO 按以下顺序依次更新各车型智能体：')

add_body(doc, '    ① 公交智能体（优先级最高）：公交车载客量大、充电需求刚性，优先保障。')
add_body(doc, '    ② 乘用车智能体（优先级中等）：兼顾效率与公平性。')
add_body(doc, '    ③ 两轮车智能体（优先级最低）：灵活调度，利用碎片时间充电。')

add_body(doc,
    '每次更新时，当前车型的策略参数以其他车型的最新策略为条件进行更新，'
    '避免多智能体协同优化中的不稳定性问题（又称为"环境非平稳性"问题）。')

add_body(doc, '通俗理解：', bold=True)
add_body(doc,
    '• 顺序更新的好处：避免了"公交想往东，乘用车想往西，两轮车想往南"这种多智能体互相拉扯、谁都学不好的情况。'
    '先固定公交智能体（优先保障），再让乘用车在公交约束下优化自己的策略，最后两轮车在前面两者的基础上找机会。')
add_body(doc,
    '• 条件更新：每个智能体更新时，都会参考其他智能体当前学到的"知识"（最新策略），'
    '而不是用旧的。这确保了每次更新都是基于最新环境状态，不会因为其他智能体策略改变而产生误导性梯度。')
add_body(doc,
    '• 对比 MAPPO：MAPPO 对所有智能体同时更新，虽然计算效率高，但容易出现"谁先学得快谁主导"的问题，'
    '低优先级车型被高优先级车型的策略带偏。HAPPO 的顺序更新有效缓解了这一问题。')

# ==================== 算法3：HAPPO主循环 ====================
add_heading(doc, '算法 3.3  HAPPO-GNN-RL 完整训练算法', level=2)

algo3 = [
    '# ── 超参数 ──',
    '# N：站点数，K：车型数（=3），L：GAT层数，H： episodes per batch',
    '# γ：折扣因子，ε：PPO裁剪系数，η：学习率',
    '',
    '# ── 主循环 ──',
    'for episode = 1 to MAX_EPISODES:',
    '',
    '    # ── [Phase 1] GAT 空间特征提取 ──',
    '    S ← 构建物理先验矩阵（互信息 + 格兰杰因果）  # 式（3.1）',
    '    H ← GAT_Feature_Extract(N, d, L, K, S)       # 算法 3.1，返回 {h_i^L}',
    '',
    '    # ── [Phase 2] 环境交互，收集经验数据 ──',
    '    for each station i = 1 to N:',
    '        s_i ← 获取当前状态（负荷、队列、绿电、电价）',
    '        h_i ← 查表获取 GAT 特征 H[i]',
    '        π_norm[i], a_i^* ← MaskedPolicy(h_i, M(s_i))   # 算法 3.2',
    '        r_t, s_{i,next} ← 执行动作 a_i^*，获取奖励与下一状态',
    '        存储经验 (s_i, a_i^*, r_t, s_{i,next}) 到缓冲池',
    '    end for',
    '',
    '    # ── [Phase 3] 优势函数估计 ──',
    '    for each station i:',
    '        A_i ← GAE(优势函数广义自举估计)',
    '    end for',
    '',
    '    # ── [Phase 4] HAPPO 顺序策略更新（异质车型）──',
    '    vehicle_types ← [公交车, 乘用车, 两轮车]     # 按优先级排序',
    '    for each type v in vehicle_types:',
    '        stations_v ← 属于车型 v 的站点集合',
    '        for each 策略更新轮 epoch = 1 to PPO_EPOCHS:',
    '            # 计算 PPO 裁剪损失',
    '            for each station i in stations_v:',
    '                π_old ← 当前策略 π_θ_old',
    '                π_new ← 新策略 π_θ',
    '                r_t ← π_new / π_old',
    '                L_i ← min(r_t · A_i, clip(r_t, 1-ε, 1+ε) · A_i)',
    '            end for',
    '            L_CLIP ← mean(L_i) + λ · policy_entropy      # 加入熵正则，防止过早收敛',
    '            θ_new ← θ_old - η · ∇_θ L_CLIP               # 式（3.14），Adam优化器',
    '            θ_old ← θ_new',
    '        end for',
    '        # 当前车型更新完毕，固定其策略参数，供下一车型使用',
    '        π_{v,fixed} ← θ_new',
    '    end for',
    '',
    '    # ── [Phase 5] 评估与记录 ──',
    '    if episode % EVAL_INTERVAL == 0:',
    '        R_avg ← 在测试集上运行当前策略，计算平均奖励',
    '        T_avg ← 平均等待时间',
    '        G_avg ← 平均绿电利用率',
    '        log(episode, R_avg, T_avg, G_avg)',
    '        if 收敛条件满足: break',
    '    end if',
    'end for',
]

algo3_explain = [
    '【通俗讲解】',
    '    ① Phase 1（GAT）：对全网站点做一次特征提取，得到每个站点融合了空间关系的特征向量。'
    '    物理先验（S矩阵）告诉GAT"历史上哪些站有关系"，GAT在此基础上学习"当前实际关系有多强"。',
    '    ② Phase 2（交互）：每个站点按当前策略选动作，执行后获得奖励。把整条经验轨迹存入缓冲池，供后续训练用。',
    '    ③ Phase 3（GAE）：用广义自举估计（Generalized Advantage Estimation）计算每个动作的优势函数 A_i。'
    '    GAE 在偏差与方差之间做平衡，比简单蒙特卡洛更稳定。',
    '    ④ Phase 4（顺序更新）：这是 HAPPO 的核心。先更新公交策略（固定）→ 再更新乘用车（参考公交）→'
    '    最后更新两轮车（参考前两者）。每型智能体独立算 PPO 损失、独立更新参数，互不干扰但相互协调。',
    '    ⑤ Phase 5（评估）：每50-100轮做一次评估，看奖励是否在上升、是否收敛。收敛则提前停止。',
    '    ⑥ 熵正则（λ·policy_entropy）：在损失里加一项策略熵（策略分布的混乱程度）。'
    '    熵越大 = 策略越"随机" = 探索越多。加熵项防止策略在找到局部最优后过早"锁死"，继续探索以获得更好的解。',
]

add_algo_block(doc, algo3, title="", explain_lines=algo3_explain)

# ==================== 3.7 模型总结 ====================
add_heading(doc, '3.7  模型框架总结', level=2)

add_body(doc,
    '综合以上各模块，HAPPO-GNN-RL 模型的完整工作流程如下：')

# 汇总流程图（文字版）
flow_steps = [
    'Step 1｜物理先验构建：基于互信息和格兰杰因果检验，构建站点关联强度矩阵 $S$（离线计算）',
    'Step 2｜GAT 空间编码：输入站点原始特征 $H^{(0)}$，通过 $L$ 层 GAT（含多头注意力、残差连接）输出空间感知特征 $H^{(L)}$',
    'Step 3｜策略网络推理：对每个站点，用 MLP 将特征 $h_i$ 映射为动作概率，施加动作掩码后输出最优动作 $a_i^*$',
    'Step 4｜环境交互：将动作 $a_i^*$ 施加于充电站环境，获得奖励 $R_t$ 和下一状态 $s_{t+1}$',
    'Step 5｜HAPPO 训练：按车型优先级顺序更新各策略网络，用 PPO 裁剪损失优化策略参数 $\\theta$',
    'Step 6｜循环迭代：重复 Step 1-5，直到策略收敛或达到最大训练轮数',
]

for i, step in enumerate(flow_steps, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(step)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_body(doc,
    '模型实现了从原始数据到调度决策的端到端优化，无需人工设计调度规则。'
    '物理先验提供了领域知识引导，GAT 捕捉了实时空间依赖，HAPPO 实现了多车型差异化协同调度。'
    '三者协同，构成了完整的充电站动态调度优化系统。')

# ==================== 公式汇总表 ====================
add_heading(doc, '3.8  本章公式汇总', level=2)

# 创建汇总表格
table = doc.add_table(rows=15, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['编号', '公式名称', '核心表达式']
header_row = table.rows[0]
for j, h in enumerate(headers):
    cell = header_row.cells[j]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(255, 255, 255)
    # 设置表头背景色
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2C5F8A')
    tcPr.append(shd)

formulas_summary = [
    ('式(3.1)', '站点关联强度矩阵', '$S_{ij} = MI(i,j) + Granger(i\\to j)$'),
    ('式(3.2)', '动态注意力分数', '$\\omega_{ij} = LeakyReLU(W_a · [h_i \\oplus h_j])$'),
    ('式(3.3)', '注意力融合', '$e_{ij} = \\omega_{ij} + \\gamma · S_{ij}$'),
    ('式(3.4)', '注意力归一化', '$\\alpha_{ij} = softmax(e_{ij})$'),
    ('式(3.5)', '图特征聚合', "$h_i' = \\sum_{j\\in\\mathcal{N}_i} \\alpha_{ij}·W_V h_j$"),
    ('式(3.6)', '多头注意力聚合', '$h_i^{(l+1)} = \\bigoplus_{k=1}^K \\sigma(...)$'),
    ('式(3.7)', '残差连接', '$h_i^{(l+1)} = \\sigma(h_i\') + h_i^{(l)}$'),
    ('式(3.8)', '策略网络 MLP', '$z_1 = W_1 h_i + b_1; \\ z_3 = W_2 · LeakyReLU(z_1) + b_2$'),
    ('式(3.9)', '动作掩码', '$\\pi_{masked} = \\pi_\\theta \\odot M(s_i)$'),
    ('式(3.10)', '掩码重归一化', '$\\pi_{norm} = \\pi_{masked} / \\sum\\pi_{masked}$'),
    ('式(3.11)', '最终调度动作', '$a_i^* = \\arg\\max_{a_i} \\pi_{norm}(a_i|s_i)$'),
    ('式(3.12)', '综合奖励函数', '$R_t = -(w_1T_t + w_2P_t + w_3L_t - w_4G_t)$'),
    ('式(3.13)', 'PPO 裁剪损失', '$L^{CLIP} = E_t[\\min(r_t A_t, clip(r_t,1-\\epsilon,1+\\epsilon) A_t)]$'),
    ('式(3.14)', '梯度更新', '$\\theta_{new} = \\theta_{old} - \\eta · \\nabla_\\theta L^{CLIP}$'),
]

for i, (num, name, expr) in enumerate(formulas_summary):
    row = table.rows[i + 1]
    row.cells[0].text = num
    row.cells[1].text = name
    row.cells[2].text = expr
    for cell in row.cells:
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(cell.text)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        # 斑马色
        if i % 2 == 0:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'EEF4FA')
            tcPr.append(shd)

add_table_caption(doc, '表 3.1  HAPPO-GNN-RL 模型公式汇总')

# ==================== 保存 ====================
output_path = r'C:\Users\罗\WorkBuddy\2026-05-04-task-5\第3章_模型建立_算法章节.docx'
doc.save(output_path)
print(f'✅ 算法章节已保存至: {output_path}')
print(f'   共包含 3 个算法 + 14 个公式 + 详细通俗讲解')
